"""Newsletter renderer.

Takes a structured content dict + an optional client config and produces a
styled .docx. Defaults preserve Jacob's Gazette branding so the original cron
keeps working without flags.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    import yaml  # PyYAML — only required if --config is passed
except ImportError:
    yaml = None

NAVY = RGBColor(0x0A, 0x1F, 0x3D)
RED = RGBColor(0xC8, 0x10, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_GRAY = RGBColor(0x22, 0x22, 0x22)
MUTED_GRAY = RGBColor(0x66, 0x66, 0x66)

ASSETS = Path(__file__).parent / "assets"

# Branding — overridable via apply_config(). Globals so the existing helper
# functions (which reference NAVY/RED directly) don't need rewiring.
TITLE_TEXT = "JACOB'S GAZETTE"
SUBTITLE_TEXT = "A WEEKLY BRIEF"
FOOTER_TEXT = "Jacob's Gazette"
DEFAULT_TAGLINE = "Faith • Tech • Tennessee • Iron"


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hex(color: RGBColor) -> str:
    """RGBColor → 6-char uppercase hex (no #), for OOXML attributes."""
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def apply_config(config: dict | None) -> None:
    """Override branding globals from a loaded config.yaml dict."""
    global NAVY, RED, ASSETS, TITLE_TEXT, SUBTITLE_TEXT, FOOTER_TEXT, DEFAULT_TAGLINE
    if not config:
        return
    nl = config.get("newsletter", {}) or {}
    br = config.get("branding", {}) or {}
    cl = config.get("client", {}) or {}
    if nl.get("title"):
        TITLE_TEXT = nl["title"]
        FOOTER_TEXT = nl["title"].title().replace("'S", "'s")
    if nl.get("subtitle"):
        SUBTITLE_TEXT = nl["subtitle"]
    if nl.get("tagline"):
        DEFAULT_TAGLINE = nl["tagline"]
    if br.get("primary_color"):
        NAVY = _hex_to_rgb(br["primary_color"])
    if br.get("accent_color"):
        RED = _hex_to_rgb(br["accent_color"])
    # Per-client assets dir takes priority; fall back to default if missing.
    slug = cl.get("slug")
    if slug:
        candidate = Path(__file__).parent / "clients" / slug / "assets"
        if candidate.exists():
            ASSETS = candidate


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_paragraph_border(paragraph, *, position: str, color_hex: str, sz: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    bd = OxmlElement(f"w:{position}")
    bd.set(qn("w:val"), "single")
    bd.set(qn("w:sz"), str(sz))
    bd.set(qn("w:space"), "1")
    bd.set(qn("w:color"), color_hex)
    pbdr.append(bd)


def _styled_run(paragraph, text: str, *, bold=False, italic=False, size=12, color=BODY_GRAY, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def _add_hyperlink(paragraph, url: str, text: str, *, size=11, color=NAVY, italic=False) -> None:
    """python-docx doesn't expose hyperlinks directly — drop in the raw XML."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    r_pr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    r_pr.append(color_el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    if italic:
        r_pr.append(OxmlElement("w:i"))
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _section_heading(doc: Document, title: str, kicker: str | None = None) -> None:
    """A red kicker line + navy section title with red underline.
    Consumes a pending page break (set by build_newsletter between sections) by
    inlining it as a run, which avoids the orphan-paragraph blank-page bug from
    doc.add_page_break()."""
    pending_break = getattr(doc, "_jg_section_break_pending", False)

    if kicker:
        kp = doc.add_paragraph()
        if pending_break:
            kp.add_run().add_break(WD_BREAK.PAGE)
            doc._jg_section_break_pending = False
            pending_break = False
        kp.paragraph_format.space_before = Pt(14)
        kp.paragraph_format.space_after = Pt(0)
        _styled_run(kp, kicker.upper(), bold=True, size=9, color=RED, font="Calibri")

    p = doc.add_paragraph()
    if pending_break:
        p.add_run().add_break(WD_BREAK.PAGE)
        doc._jg_section_break_pending = False
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    _styled_run(p, title, bold=True, size=16, color=NAVY, font="Calibri")
    _set_paragraph_border(p, position="bottom", color_hex=_hex(RED), sz=12)


def _body_paragraph(doc: Document, text: str, *, italic=False, size=12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    _styled_run(p, text, size=size, color=BODY_GRAY, italic=italic)


def _bullet(doc: Document, text: str, *, size=12) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = BODY_GRAY
        run.font.name = "Calibri"
    if not p.runs:
        _styled_run(p, text, size=size)
    else:
        p.runs[0].text = text


def _maybe_image(doc: Document, image_path: str | None, *, width_in: float = 4.5, max_height_in: float | None = None, caption: str | None = None) -> None:
    """Insert a centered image with optional caption. Silently skips if path missing.
    If max_height_in is set, the image is constrained by height instead of width so
    portrait-ratio photos never overflow the page."""
    if not image_path:
        return
    ip = Path(image_path)
    if not ip.is_absolute():
        ip = Path(__file__).parent / image_path
    if not ip.exists():
        return
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(2)
    if max_height_in is not None:
        pp.add_run().add_picture(str(ip), height=Inches(max_height_in))
    else:
        pp.add_run().add_picture(str(ip), width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        _styled_run(cp, caption, italic=True, size=9, color=MUTED_GRAY)


def _source_line(doc: Document, source_text: str, source_url: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    _styled_run(p, "Source: ", italic=True, size=10, color=MUTED_GRAY)
    if source_url:
        _add_hyperlink(p, source_url, source_text, size=10, color=NAVY, italic=True)
    else:
        _styled_run(p, source_text, italic=True, size=10, color=MUTED_GRAY)


# ---------- Section renderers ----------


def render_article(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    paras = section.get("paragraphs") or section.get("body", [])
    for para in paras:
        if isinstance(para, str) and para.startswith("Source:"):
            raw = para[7:].strip()
            if raw.startswith("http"):
                _source_line(doc, raw, raw)
            else:
                _source_line(doc, raw)
        else:
            _body_paragraph(doc, para)
    if section.get("source"):
        src = section["source"]
        _source_line(doc, src.get("text", "Read more"), src.get("url"))


def render_events_list(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])
    _maybe_image(doc, section.get("image_path"), width_in=4.5, caption=section.get("image_caption"))
    for item in section.get("items", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        # Strip the auto-generated run from the bullet style
        for r in list(p.runs):
            r.text = ""
        if item.get("url"):
            _add_hyperlink(p, item["url"], item["name"], size=11, color=NAVY)
            # Force the linked name to render bold
            link_el = p._p.findall(qn("w:hyperlink"))[-1]
            r_pr = link_el.find(".//" + qn("w:rPr"))
            if r_pr is not None and r_pr.find(qn("w:b")) is None:
                r_pr.append(OxmlElement("w:b"))
        else:
            _styled_run(p, item["name"], bold=True, size=11)
        _styled_run(p, f" — {item['date']}", size=11)
        if item.get("location"):
            _styled_run(p, f" • {item['location']}", size=11, color=MUTED_GRAY)


def render_strava(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    stats = section.get("stats", {})
    if stats:
        table = doc.add_table(rows=2, cols=4)
        table.autofit = False
        labels = ["Runs", "Miles", "Time", "Avg Pace"]
        values = [
            str(stats.get("runs") or stats.get("activities", "—")),
            str(stats.get("total_miles") or stats.get("distance_miles", "—")),
            stats.get("total_time") or stats.get("moving_time", "—"),
            stats.get("avg_pace", "—"),
        ]
        for i, lbl in enumerate(labels):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(cell, _hex(NAVY))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            _styled_run(p, lbl.upper(), bold=True, size=9, color=WHITE)
        for i, val in enumerate(values):
            cell = table.rows[1].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            _styled_run(p, str(val), bold=True, size=14, color=NAVY)
    if section.get("summary"):
        _body_paragraph(doc, section["summary"])
    activities = section.get("activities", [])
    if activities:
        for a in activities:
            p = doc.add_paragraph(style="List Bullet")
            for r in list(p.runs):
                r.text = ""
            _styled_run(p, a["name"], bold=True, size=11)
            dist = a.get("distance") or (f"{a['miles']} mi" if "miles" in a else "")
            detail = f" — {a.get('date','')} • {dist} • {a.get('time','')} • {a.get('pace','')}"
            _styled_run(p, detail, size=11, color=MUTED_GRAY)


def render_devotional(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    # Verse block
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _set_paragraph_border(p, position="left", color_hex=_hex(RED), sz=18)
    _styled_run(p, f"“{section['verse_text']}”", italic=True, size=12, color=BODY_GRAY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25)
    p2.paragraph_format.space_after = Pt(6)
    _styled_run(p2, f"— {section['verse_ref']}", bold=True, size=11, color=NAVY)
    if section.get("reflection"):
        _body_paragraph(doc, section["reflection"])


def render_chess(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _styled_run(p, "Opening of the week: ", bold=True, size=12, color=NAVY)
    _styled_run(p, section["opening_name"], bold=True, size=12, color=RED)
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])
    _maybe_image(doc, section.get("image_path"), width_in=2.6, caption=section.get("image_caption"))
    if section.get("key_ideas"):
        for idea in section["key_ideas"]:
            _bullet(doc, idea)
    if section.get("videos"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        _styled_run(p, "Study videos:", bold=True, size=11, color=NAVY)
        for v in section["videos"]:
            vp = doc.add_paragraph(style="List Bullet")
            for r in list(vp.runs):
                r.text = ""
            _add_hyperlink(vp, v["url"], v["title"], size=11, color=NAVY)
            if v.get("channel"):
                _styled_run(vp, f" — {v['channel']}", size=10, color=MUTED_GRAY, italic=True)


def render_recipe(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _styled_run(p, section["recipe_name"], bold=True, size=13, color=NAVY)
    meta_bits = []
    if section.get("servings"):
        meta_bits.append(f"Serves {section['servings']}")
    if section.get("time"):
        meta_bits.append(section["time"])
    if section.get("calories"):
        meta_bits.append(section["calories"])
    if meta_bits:
        mp = doc.add_paragraph()
        mp.paragraph_format.space_after = Pt(4)
        _styled_run(mp, " • ".join(meta_bits), italic=True, size=10, color=MUTED_GRAY)
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])
    _maybe_image(doc, section.get("image_path"), max_height_in=1.5, caption=section.get("image_caption"))

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    # Ingredients
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    _styled_run(lp, "INGREDIENTS", bold=True, size=10, color=RED)
    for ing in section.get("ingredients", []):
        ip = left.add_paragraph()
        ip.paragraph_format.space_after = Pt(0)
        _styled_run(ip, f"• {ing}", size=11)
    # Directions
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(2)
    _styled_run(rp, "DIRECTIONS", bold=True, size=10, color=RED)
    for i, step in enumerate(section.get("directions", []), 1):
        sp = right.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        _styled_run(sp, f"{i}. ", bold=True, size=11, color=NAVY)
        _styled_run(sp, step, size=11)


def render_lifehack(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _styled_run(p, section["hack_name"], bold=True, size=13, color=NAVY)
    for para in (section.get("paragraphs") or section.get("body", [])):
        _body_paragraph(doc, para)


def render_flights(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])
    items = section.get("items", [])
    if not items:
        _body_paragraph(doc, section.get("empty_message", "No abnormally low fares found this week — checking again next Monday."), italic=True)
        return
    # Free-form format: route / window / deal / tip
    if "route" in items[0]:
        for item in items:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_before = Pt(6)
            rp.paragraph_format.space_after = Pt(1)
            _styled_run(rp, item.get("route", ""), bold=True, size=12, color=NAVY)
            if item.get("window"):
                wp = doc.add_paragraph()
                wp.paragraph_format.space_after = Pt(1)
                _styled_run(wp, item["window"], italic=True, size=10, color=MUTED_GRAY)
            if item.get("deal"):
                _body_paragraph(doc, item["deal"])
            if item.get("tip"):
                tp = doc.add_paragraph()
                tp.paragraph_format.space_after = Pt(4)
                _styled_run(tp, "Tip: ", bold=True, size=10, color=NAVY)
                _styled_run(tp, item["tip"], size=10, color=MUTED_GRAY)
        if section.get("note"):
            np = doc.add_paragraph()
            np.paragraph_format.space_before = Pt(4)
            _styled_run(np, section["note"], italic=True, size=10, color=MUTED_GRAY)
        return
    # Structured table format: origin / destination / depart / return_date / price / carrier
    table = doc.add_table(rows=len(items) + 1, cols=5)
    table.autofit = True
    headers = ["Route", "Depart", "Return", "Price", "Carrier"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, _hex(NAVY))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        _styled_run(p, h.upper(), bold=True, size=9, color=WHITE)
    for r, item in enumerate(items, 1):
        cells = table.rows[r].cells
        route_p = cells[0].paragraphs[0]
        route_p.paragraph_format.space_after = Pt(0)
        if item.get("link"):
            _add_hyperlink(route_p, item["link"], f"{item.get('origin','RDU')} → {item.get('destination','—')}", size=11, color=NAVY)
        else:
            _styled_run(route_p, f"{item.get('origin','RDU')} → {item.get('destination','—')}", bold=True, size=11)
        for i, key in enumerate(["depart", "return_date", "price", "carrier"], start=1):
            cp = cells[i].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            _styled_run(cp, str(item.get(key, "—")), size=11)
    if section.get("note"):
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(4)
        _styled_run(np, section["note"], italic=True, size=10, color=MUTED_GRAY)


def render_pentest(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    if section.get("generated_at"):
        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(4)
        _styled_run(gp, f"Scan generated: {section['generated_at']}", italic=True, size=10, color=MUTED_GRAY)
    if section.get("summary"):
        _body_paragraph(doc, section["summary"])

    devices = section.get("network_map", [])
    if devices:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(4)
        sub.paragraph_format.space_after = Pt(2)
        _styled_run(sub, "Network Map", bold=True, size=11, color=NAVY)
        table = doc.add_table(rows=len(devices) + 1, cols=4)
        table.autofit = True
        for i, h in enumerate(["IP", "Device", "Open Ports", "Notes"]):
            cell = table.rows[0].cells[i]
            _set_cell_shading(cell, _hex(NAVY))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _styled_run(p, h.upper(), bold=True, size=9, color=WHITE)
        for r, d in enumerate(devices, 1):
            cells = table.rows[r].cells
            for idx, key in enumerate(["ip", "label", "open_ports", "notes"]):
                cp = cells[idx].paragraphs[0]
                cp.paragraph_format.space_after = Pt(0)
                val = d.get(key, "—")
                if isinstance(val, list):
                    val = ", ".join(str(v) for v in val) if val else "—"
                _styled_run(cp, str(val), size=10)

    findings = section.get("key_findings", [])
    if findings:
        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(6)
        fp.paragraph_format.space_after = Pt(2)
        _styled_run(fp, "Key Findings", bold=True, size=11, color=RED)
        for finding in findings:
            _bullet(doc, finding, size=11)

    recs = section.get("recommendations", [])
    if recs:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after = Pt(2)
        _styled_run(rp, "Recommendations", bold=True, size=11, color=NAVY)
        for rec in recs:
            _bullet(doc, rec, size=11)


def render_crossword(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])

    image_path = section.get("image_path")
    if image_path:
        ip = Path(image_path)
        if not ip.is_absolute():
            ip = Path(__file__).parent / image_path
        if ip.exists():
            pp = doc.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(2)
            pp.paragraph_format.space_after = Pt(4)
            pp.add_run().add_picture(str(ip), width=Inches(3.6))

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    _styled_run(lp, "ACROSS", bold=True, size=10, color=RED)
    for clue in section.get("across", []):
        cp = left.add_paragraph()
        cp.paragraph_format.space_after = Pt(1)
        _styled_run(cp, f"{clue.get('number','')}. ", bold=True, size=10, color=NAVY)
        _styled_run(cp, clue.get("clue", ""), size=10)

    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(2)
    _styled_run(rp, "DOWN", bold=True, size=10, color=RED)
    for clue in section.get("down", []):
        cp = right.add_paragraph()
        cp.paragraph_format.space_after = Pt(1)
        _styled_run(cp, f"{clue.get('number','')}. ", bold=True, size=10, color=NAVY)
        _styled_run(cp, clue.get("clue", ""), size=10)


def render_vehicle_listings(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    if section.get("intro"):
        _body_paragraph(doc, section["intro"])
    _maybe_image(doc, section.get("image_path"), width_in=3.4, caption=section.get("image_caption"))
    items = section.get("items", [])
    if not items:
        _body_paragraph(doc, section.get("empty_message", "No matches this week — will keep watching."), italic=True)
        return
    table = doc.add_table(rows=len(items) + 1, cols=5)
    table.autofit = True
    for i, h in enumerate(["Year", "Miles", "Price", "Location", "Source"]):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, _hex(NAVY))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        _styled_run(p, h.upper(), bold=True, size=9, color=WHITE)
    for r, item in enumerate(items, 1):
        cells = table.rows[r].cells
        for idx, key in enumerate(["year", "miles", "price", "location"]):
            cp = cells[idx].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            _styled_run(cp, str(item.get(key, "—")), size=11)
        sp = cells[4].paragraphs[0]
        sp.paragraph_format.space_after = Pt(0)
        if item.get("url"):
            _add_hyperlink(sp, item["url"], item.get("source", "Listing"), size=11, color=NAVY)
        else:
            _styled_run(sp, item.get("source", "—"), size=11)
    if section.get("market_note"):
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(8)
        _styled_run(np, section["market_note"], italic=True, size=10, color=MUTED_GRAY)


def render_concert_airbnb(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    concert = section.get("concert")
    if not concert:
        _body_paragraph(doc, section.get("empty_message", "No upcoming Morgan Wallen concerts within range — will keep watching."), italic=True)
        return
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(2)
    _styled_run(cp, concert.get("artist", "Concert"), bold=True, size=13, color=NAVY)
    _styled_run(cp, "  •  ", size=11)
    _styled_run(cp, f"{concert.get('date','TBD')} • {concert.get('venue','TBD')}, {concert.get('city','TBD')}", size=11, color=BODY_GRAY)
    if concert.get("ticket_url"):
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(6)
        _add_hyperlink(tp, concert["ticket_url"], "→ Tickets", size=11, color=RED)

    listings = section.get("airbnb_listings", [])
    if listings:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(2)
        sub.paragraph_format.space_after = Pt(2)
        _styled_run(sub, "Airbnb matches (2 BR · 4.6+ rating · within 10 mi)", bold=True, size=11, color=NAVY)
        for l in listings:
            lp = doc.add_paragraph(style="List Bullet")
            for r in list(lp.runs):
                r.text = ""
            if l.get("url"):
                _add_hyperlink(lp, l["url"], l.get("name", "Listing"), size=11, color=NAVY)
            else:
                _styled_run(lp, l.get("name", "Listing"), bold=True, size=11)
            detail = []
            if l.get("price_per_night"):
                detail.append(l["price_per_night"])
            if l.get("rating"):
                detail.append(f"★ {l['rating']}")
            if l.get("distance"):
                detail.append(l["distance"])
            if detail:
                _styled_run(lp, " — " + " • ".join(detail), size=10, color=MUTED_GRAY)
    else:
        _body_paragraph(doc, "No Airbnb matches found for the criteria.", italic=True)
    if section.get("notes"):
        np = doc.add_paragraph()
        np.paragraph_format.space_before = Pt(6)
        _styled_run(np, section["notes"], italic=True, size=10, color=MUTED_GRAY)


def render_alpaca_summary(doc: Document, section: dict) -> None:
    _section_heading(doc, section["title"], kicker=section.get("kicker"))
    raw = section.get("stats", {})
    # Normalize both schema variants into display-ready strings
    stats: dict = {}
    eq = raw.get("portfolio_value") or raw.get("equity")
    stats["portfolio_value"] = f"${eq:,.2f}" if isinstance(eq, (int, float)) else (str(eq) if eq else "—")
    dc = raw.get("day_change", 0)
    stats["day_change"] = f"{dc:+,.2f}" if isinstance(dc, (int, float)) else str(dc)
    tr = raw.get("total_return") or raw.get("total_return_pct")
    stats["total_return"] = f"{tr:+.2f}%" if isinstance(tr, (int, float)) else (str(tr) if tr else "—")
    cash = raw.get("cash")
    stats["cash"] = f"${cash:,.2f}" if isinstance(cash, (int, float)) else (str(cash) if cash else "—")
    if stats:
        table = doc.add_table(rows=2, cols=4)
        table.autofit = False
        labels = ["Portfolio", "Day Change", "Total Return", "Cash"]
        keys = ["portfolio_value", "day_change", "total_return", "cash"]
        for i, lbl in enumerate(labels):
            cell = table.rows[0].cells[i]
            _set_cell_shading(cell, _hex(NAVY))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            _styled_run(p, lbl.upper(), bold=True, size=9, color=WHITE)
        for i, key in enumerate(keys):
            cell = table.rows[1].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            val = stats.get(key, "—")
            color = NAVY
            if key in ("day_change", "total_return") and isinstance(val, str):
                if val.strip().startswith("-"):
                    color = RGBColor(0xB0, 0x10, 0x10)
                elif val.strip().startswith("+") or any(c.isdigit() for c in val):
                    color = RGBColor(0x0A, 0x6B, 0x2C)
            _styled_run(p, str(val), bold=True, size=14, color=color)
    if section.get("summary"):
        _body_paragraph(doc, section["summary"])
    holdings = section.get("top_holdings", [])
    if holdings:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(4)
        hp.paragraph_format.space_after = Pt(2)
        _styled_run(hp, "Top Holdings", bold=True, size=11, color=NAVY)
        table = doc.add_table(rows=len(holdings) + 1, cols=4)
        table.autofit = True
        for i, h in enumerate(["Symbol", "Qty", "Avg Cost", "Market Value"]):
            cell = table.rows[0].cells[i]
            _set_cell_shading(cell, _hex(NAVY))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _styled_run(p, h.upper(), bold=True, size=9, color=WHITE)
        for r, h in enumerate(holdings, 1):
            cells = table.rows[r].cells
            for idx, key in enumerate(["symbol", "qty", "avg_cost", "market_value"]):
                cp = cells[idx].paragraphs[0]
                cp.paragraph_format.space_after = Pt(0)
                _styled_run(cp, str(h.get(key, "—")), size=11, bold=(idx == 0))
    activity = section.get("recent_activity", [])
    if activity:
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(4)
        ap.paragraph_format.space_after = Pt(2)
        _styled_run(ap, "Recent Activity (last 7 days)", bold=True, size=11, color=NAVY)
        for a in activity:
            _bullet(doc, a, size=11)


# ---------- Title page + TOC ----------


def render_title_page(doc: Document, content: dict) -> None:
    """Render the cover page. Uses a full-bleed image if title_page.image_path is set,
    otherwise falls back to a styled text-only cover."""
    title_meta = content.get("title_page") or {}
    image_path = title_meta.get("image_path")
    if image_path:
        ip = Path(image_path)
        if not ip.is_absolute():
            ip = ASSETS.parent / image_path
        if ip.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            # Size by height so any aspect ratio fits within the page content area.
            # Letter is 11" tall; with 1.4" top + 0.8" bottom margins, usable ≈ 8.8".
            run = p.add_run()
            run.add_picture(str(ip), height=Inches(8.4))
            # Page break in the SAME paragraph — avoids a stray empty page from
            # add_page_break() creating its own paragraph after the cover image.
            run.add_break(WD_BREAK.PAGE)
            return
    # Text fallback
    for _ in range(4):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    _styled_run(title_p, TITLE_TEXT, bold=True, size=44, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(20)
    _styled_run(sub, content.get("issue_label", "").upper(), bold=True, size=14, color=RED)
    if title_meta.get("tagline"):
        tag = doc.add_paragraph()
        tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(tag, title_meta["tagline"], italic=True, size=14, color=BODY_GRAY)
    # Force a page break after the text-only cover so the TOC starts fresh on page 2
    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(0)
    pb.add_run().add_break(WD_BREAK.PAGE)


def render_toc(doc: Document, content: dict) -> None:
    """Static table of contents. Page numbering assumes title page = 1, TOC = 2,
    sections start at page 3 with one section per page."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(10)
    _styled_run(h, "Table of Contents", bold=True, size=26, color=NAVY)
    _set_paragraph_border(h, position="bottom", color_hex=_hex(RED), sz=18)

    sections = content.get("sections", [])
    for idx, s in enumerate(sections):
        page_num = idx + 3
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(4)
        line.paragraph_format.line_spacing = 1.4
        # Right-aligned tab stop with dot leader for the page number
        ts = line.paragraph_format.tab_stops
        ts.add_tab_stop(Inches(6.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        _styled_run(line, f"{idx + 1:>2}. ", bold=True, size=12, color=RED)
        _styled_run(line, s.get("title", "Untitled"), size=12, color=BODY_GRAY)
        line.add_run("\t")
        _styled_run(line, str(page_num), bold=True, size=12, color=NAVY)


# ---------- Header / footer / chrome ----------


def _build_header(doc: Document, issue_label: str) -> None:
    section = doc.sections[0]
    # Title page (page 1) gets no header/footer; all later pages get the masthead.
    section.different_first_page_header_footer = True
    fp_h = section.first_page_header
    fp_h.is_linked_to_previous = False
    fp_f = section.first_page_footer
    fp_f.is_linked_to_previous = False

    header = section.header
    header.is_linked_to_previous = False

    # Logo paragraph
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    logo_path = ASSETS / "header_logo.png"
    if logo_path.exists():
        run.add_picture(str(logo_path), width=Inches(6.5))

    # Issue label / tagline below the logo
    tagline = header.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(0)
    _styled_run(
        tagline,
        f"{SUBTITLE_TEXT}  •  {issue_label.upper()}",
        bold=True,
        size=9,
        color=NAVY,
    )
    _set_paragraph_border(tagline, position="bottom", color_hex=_hex(RED), sz=18)


def _build_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, f"{FOOTER_TEXT}  •  Page ", size=9, color=MUTED_GRAY)

    # Page number field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_GRAY
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.4)  # leaves room for header
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)


def _set_page_borders(doc: Document, color_hex: str | None = None, sz: int = 18) -> None:
    """Add a uniform border around every page. sz is in eighths of a point
    (18 ≈ 2.25pt). offsetFrom='page' anchors the border to the page edge."""
    if color_hex is None:
        color_hex = _hex(NAVY)
    for section in doc.sections:
        sect_pr = section._sectPr
        # Remove any existing pgBorders so this is the source of truth
        for existing in sect_pr.findall(qn("w:pgBorders")):
            sect_pr.remove(existing)
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        for side in ("top", "left", "bottom", "right"):
            bd = OxmlElement(f"w:{side}")
            bd.set(qn("w:val"), "single")
            bd.set(qn("w:sz"), str(sz))
            bd.set(qn("w:space"), "24")
            bd.set(qn("w:color"), color_hex)
            pg_borders.append(bd)
        sect_pr.append(pg_borders)


# ---------- Main ----------


RENDERERS = {
    "article": render_article,
    "events_list": render_events_list,
    "strava": render_strava,
    "devotional": render_devotional,
    "chess": render_chess,
    "chess_opening": render_chess,
    "recipe": render_recipe,
    "lifehack": render_lifehack,
    "flights": render_flights,
    "pentest": render_pentest,
    "crossword": render_crossword,
    "vehicle_listings": render_vehicle_listings,
    "concert_airbnb": render_concert_airbnb,
    "alpaca_summary": render_alpaca_summary,
    "alpaca_paper": render_alpaca_summary,
}


def build_newsletter(content: dict, output_path: Path) -> None:
    doc = Document()

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)
    style.font.color.rgb = BODY_GRAY

    _set_margins(doc)
    _set_page_borders(doc)
    _build_header(doc, content.get("issue_label", ""))
    _build_footer(doc)

    # Page 1: title page (no header/footer — first-page-different is set in _build_header)
    # render_title_page emits its own trailing page break inside the image paragraph,
    # so don't add another one here (it would create a blank page 2).
    render_title_page(doc, content)

    # Page 2: table of contents
    render_toc(doc, content)

    # Pages 3+: one section per page. The first section needs a leading break;
    # subsequent breaks are emitted by _section_heading via the pending-break flag.
    doc._jg_section_break_pending = True
    for section in content.get("sections", []):
        renderer = RENDERERS.get(section["type"])
        if renderer:
            renderer(doc, section)
        else:
            _section_heading(doc, section.get("title", "Untitled"))
            _body_paragraph(doc, f"[Unknown section type: {section['type']}]")
        # Mark a page break pending for the next section's heading to consume.
        doc._jg_section_break_pending = True

    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Render newsletter .docx from content + optional client config")
    parser.add_argument("content", help="Path to content.json")
    parser.add_argument("output", help="Path to output .docx")
    parser.add_argument("--config", help="Path to client config.yaml (optional; defaults to Jacob's Gazette branding)")
    args = parser.parse_args()

    config = None
    if args.config:
        if yaml is None:
            print("ERROR: --config requires PyYAML. Run: pip install pyyaml", file=sys.stderr)
            sys.exit(2)
        config = yaml.safe_load(Path(args.config).read_text())
    apply_config(config)

    content = json.loads(Path(args.content).read_text())
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_newsletter(content, output_path)
    print(f"Wrote {output_path} (title={TITLE_TEXT}, primary=#{_hex(NAVY)}, accent=#{_hex(RED)})")


if __name__ == "__main__":
    main()
